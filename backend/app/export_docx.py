"""
BRD Quality Assessment – Word (DOCX) Report

Generates a read-only BRD view where each line of the BRD is visible and
lines with findings are highlighted with comments shown beside them.
"""

from __future__ import annotations

import io
from datetime import datetime
import sqlite3
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _get_latest_brd(conn: sqlite3.Connection):
    row = conn.execute(
        """
        SELECT doc_id, filename, full_text, line_count, upload_timestamp
        FROM documents
        WHERE doc_type = 'output_brd'
        ORDER BY doc_id DESC
        LIMIT 1
        """
    ).fetchone()
    if not row:
        return None
    return {
        "doc_id": row[0],
        "filename": row[1],
        "full_text": row[2] or "",
        "line_count": row[3] or 0,
        "upload_timestamp": row[4],
    }


def _get_findings_for_brd(conn: sqlite3.Connection, brd_doc_id: int) -> Dict[int, List[dict]]:
    cur = conn.execute(
        """
        SELECT f.finding_id,
               f.error_type,
               f.severity,
               f.line_number,
               f.description,
               f.source_reference,
               f.detected_timestamp
        FROM findings f
        JOIN chunks c ON f.chunk_id = c.chunk_id
        WHERE c.doc_id = ?
          AND f.line_number IS NOT NULL
        ORDER BY f.line_number, f.severity DESC
        """,
        (brd_doc_id,),
    )
    cols = [d[0] for d in cur.description]
    by_line: Dict[int, List[dict]] = {}
    for row in cur.fetchall():
        rec = dict(zip(cols, row))
        ln = int(rec["line_number"])
        by_line.setdefault(ln, []).append(rec)
    return by_line


def _style_document(doc: Document, filename: str, line_count: int):
    styles = doc.styles
    if "Normal" in styles:
        normal = styles["Normal"]
        if normal.font:
            normal.font.name = "Calibri"
            normal.font.size = Pt(10)

    title = doc.add_heading("BRD Quality Review", level=1)
    title.alignment = 0

    meta = doc.add_paragraph()
    meta.add_run("BRD file: ").bold = True
    meta.add_run(filename or "N/A")
    meta.add_run("\nGenerated: ").bold = True
    meta.add_run(datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))
    meta.add_run("\nTotal lines: ").bold = True
    meta.add_run(str(line_count))

    legend = doc.add_paragraph()
    legend.add_run(
        "\nHighlighted rows indicate BRD lines with one or more findings. "
        "Comments explain what is wrong with that line in business terms."
    )


def _add_table(
    doc: Document,
    lines: List[str],
    findings_by_line: Dict[int, List[dict]],
):
    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Line #"
    hdr_cells[1].text = "BRD Text"
    hdr_cells[2].text = "Findings / Review Comments"

    # Header styling
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F2937")  # dark slate
        tc_pr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = None

    for i, raw in enumerate(lines, start=1):
        line_text = raw or ""
        row = table.add_row()
        c_line, c_text, c_comments = row.cells

        c_line.text = str(i)
        c_text.text = line_text

        items = findings_by_line.get(i, [])
        if not items:
            continue

        # Highlight BRD text cell and build rich comments.
        for p in c_text.paragraphs:
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        comment_par = c_comments.paragraphs[0]
        comment_par.text = ""

        for idx, f in enumerate(items, start=1):
            if idx > 1:
                comment_par.add_run("\n")

            sev = (f.get("severity") or "").upper()
            etype = (f.get("error_type") or "").replace("_", " ")
            desc = f.get("source_reference") or f.get("description") or ""

            bullet = comment_par.add_run(f"• [{sev}] {etype}: ")
            bullet.bold = True

            body = comment_par.add_run(desc)
            body.font.size = Pt(9)

    # Column widths for better readability
    table.columns[0].width = Inches(0.7)
    table.columns[1].width = Inches(3.5)
    table.columns[2].width = Inches(3.0)


def generate_docx_report(db_path: str) -> bytes:
    """
    Build a DOCX report containing the full BRD text with
    problematic lines highlighted and comments beside them.
    """
    conn = sqlite3.connect(db_path)
    try:
        brd = _get_latest_brd(conn)
        if not brd:
            raise RuntimeError("No BRD document found in database")

        findings_by_line = _get_findings_for_brd(conn, brd["doc_id"])
        lines = (brd["full_text"] or "").splitlines()

        doc = Document()
        _style_document(doc, brd["filename"], len(lines))
        _add_table(doc, lines, findings_by_line)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    finally:
        conn.close()

