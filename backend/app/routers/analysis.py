from typing import Optional, Dict

from fastapi import APIRouter, UploadFile, File, HTTPException, status
from pydantic import BaseModel
from ..database import get_connection
from pypdf import PdfReader
import docx

from ..models import (
    insert_document,
    create_brd_chunks,
    get_document,
    get_chunks_for_brd,
    get_findings_for_brd,
    create_rule,
)

from ..detectors import (
    different_data,
    incomplete_data,
    hallucination,
    depth_mismatch,
    duplicate_data,
)

router = APIRouter(prefix="/api", tags=["analysis"])


# ─────────────────────────────────────────────
# Text extraction
# ─────────────────────────────────────────────

def extract_text_from_pdf(file: UploadFile) -> str:
    try:
        reader = PdfReader(file.file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to read PDF: {e}",
        )


def extract_text_from_docx(file: UploadFile) -> str:
    try:
        document = docx.Document(file.file)

        lines = [p.text for p in document.paragraphs]

        for table in document.tables:
            for row in table.rows:
                lines.append(
                    " | ".join(c.text.replace("\n", " ").strip() for c in row.cells)
                )

        return "\n".join(lines)

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to read DOCX: {e}",
        )


def extract_text(file: UploadFile) -> str:
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file)

    if filename.endswith(".docx"):
        return extract_text_from_docx(file)

    if filename.endswith(".txt"):
        return file.file.read().decode("utf-8", errors="ignore")

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Unsupported file type. Upload PDF, DOCX, or TXT.",
    )


# ─────────────────────────────────────────────
# Upload endpoint
# ─────────────────────────────────────────────

@router.post("/upload-documents")
async def upload_documents(
    input_sow: Optional[UploadFile] = File(None),
    input_mom: Optional[UploadFile] = File(None),
    output_brd: UploadFile = File(...),
):

    if output_brd.size == 0:
        raise HTTPException(400, "Output BRD file is empty")

    sow_doc_id = None
    mom_doc_id = None

    sow_line_count = None
    mom_line_count = None

    # SOW
    if input_sow:
        sow_text = extract_text(input_sow)

        sow_doc_id, sow_line_count = insert_document(
            "input_sow", input_sow.filename, sow_text
        )

    # MoM
    if input_mom:
        mom_text = extract_text(input_mom)

        mom_doc_id, mom_line_count = insert_document(
            "input_mom", input_mom.filename, mom_text
        )

    # BRD
    brd_text = extract_text(output_brd)

    if not brd_text.strip():
        raise HTTPException(400, "BRD file is empty")

    brd_doc_id, brd_line_count = insert_document(
        "output_brd", output_brd.filename, brd_text
    )

    chunks_created = create_brd_chunks(brd_doc_id, brd_text, chunk_size=50)

    print("UPLOAD COMPLETE")
    print("SOW lines:", sow_line_count)
    print("MoM lines:", mom_line_count)
    print("BRD lines:", brd_line_count)
    print("Chunks:", chunks_created)

    return {
        "message": "Documents uploaded successfully",
        "sow": {"doc_id": sow_doc_id, "line_count": sow_line_count},
        "mom": {"doc_id": mom_doc_id, "line_count": mom_line_count},
        "brd": {
            "doc_id": brd_doc_id,
            "line_count": brd_line_count,
            "chunks_created": chunks_created,
        },
    }


# ─────────────────────────────────────────────
# Analysis request schema
# ─────────────────────────────────────────────

class AnalysisRequest(BaseModel):
    sow_doc_id: int
    mom_doc_id: Optional[int] = None
    brd_doc_id: int


# ─────────────────────────────────────────────
# Load texts
# ─────────────────────────────────────────────

def _load_texts(payload: AnalysisRequest):

    sow_doc = get_document(payload.sow_doc_id)
    brd_doc = get_document(payload.brd_doc_id)

    if not sow_doc or not brd_doc:
        raise HTTPException(400, "Invalid document id")

    sow_text = sow_doc["full_text"] or ""
    brd_text = brd_doc["full_text"] or ""

    mom_text = ""

    if payload.mom_doc_id:
        mom_doc = get_document(payload.mom_doc_id)
        if mom_doc:
            mom_text = mom_doc["full_text"] or ""

    chunks = get_chunks_for_brd(payload.brd_doc_id)

    print("TEXT LOAD CHECK")
    print("SOW length:", len(sow_text))
    print("MoM length:", len(mom_text))
    print("BRD length:", len(brd_text))
    print("Chunks:", len(chunks))

    return sow_text, mom_text, brd_text, chunks


# ─────────────────────────────────────────────
# Summary builder
# ─────────────────────────────────────────────

def _build_summary(brd_doc_id: int):

    findings = get_findings_for_brd(brd_doc_id)

    summary: Dict[str, int] = {}

    for f in findings:
        summary.setdefault(f["error_type"], 0)
        summary[f["error_type"]] += 1

    print("FINDINGS COUNT:", len(findings))

    return {
        "summary": summary,
        "total_findings": len(findings),
        "findings": findings,
    }


# ─────────────────────────────────────────────
# FULL ANALYSIS
# ─────────────────────────────────────────────

@router.post("/run-full-analysis")
async def run_full_analysis(payload: AnalysisRequest):

    sow_text, mom_text, brd_text, chunks = _load_texts(payload)

    conn = get_connection()
    conn.execute("DELETE FROM findings")
    conn.commit()
    conn.close()

    print("Running detectors...")

    different_data.detect(sow_text, mom_text, brd_text, chunks)

    incomplete_data.detect(sow_text, brd_text, chunks)

    hallucination.detect(sow_text, mom_text, brd_text, chunks)

    depth_mismatch.detect(sow_text, brd_text, chunks)

    duplicate_data.detect(brd_text, chunks)

    return _build_summary(payload.brd_doc_id)